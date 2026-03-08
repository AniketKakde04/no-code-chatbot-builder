import os
import sys
import smtplib
import datetime
import asyncio
import nest_asyncio
import shutil
from typing import TypedDict, List, Dict, Any, Annotated
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from email.mime.base import MIMEBase
from email import encoders
import supabase
from twilio.rest import Client
from twilio.base.exceptions import TwilioRestException
from docx import Document as DocxDocument # NEW IMPORT

from google.oauth2.credentials import Credentials
from google.auth.transport.requests import Request
import gspread

# Apply nested asyncio to allow MCP client to run inside FastAPI
nest_asyncio.apply()

from langgraph.graph import StateGraph, START, END
from langgraph.graph.message import add_messages
from langgraph.prebuilt import ToolNode, tools_condition
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain_community.tools.tavily_search import TavilySearchResults
from langchain_core.messages import HumanMessage, SystemMessage, AIMessage

# --- MCP IMPORTS ---
from mcp import ClientSession, StdioServerParameters
from mcp.client.stdio import logger, stdio_client
from langchain_mcp_adapters.tools import load_mcp_tools

# --- 1. DEFINE STATE ---
class AgentState(TypedDict):
    messages: Annotated[list, add_messages]
    metadata: Dict[str, Any]
    attachment_path: str  # Path to a file attachment (e.g., from doc_writer)

# --- 2. DEFINE GLOBAL TOOLS ---
tavily_tool = TavilySearchResults(max_results=3)

# --- 3. HELPER: MCP TOOL LOADER ---
async def get_mcp_tools(command: str, args: List[str]):
    """
    Connects to a local MCP server via Stdio and returns LangChain tools.
    """
    if sys.platform == "win32":
        if command in ["npx", "npm", "npx.cmd", "npm.cmd"]:
            args = ["/c", command] + args
            command = "cmd"

    print(f"🔌 Connecting to MCP Server: {command} {args}")
    
    server_params = StdioServerParameters(
        command=command, 
        args=args, 
        env=os.environ.copy()
    )
    
    try:
        async with stdio_client(server_params) as (read, write):
            async with ClientSession(read, write) as session:
                await session.initialize()
                tools = await load_mcp_tools(session)
                print(f"✅ Loaded {len(tools)} MCP tools: {[t.name for t in tools]}")
                return tools
    except Exception as e:
        print(f"❌ MCP Connection Failed: {e}")
        return []

# --- 4. NODE FACTORIES ---

def get_llm_node(system_instruction: str, user_template: str, bind_tools: bool = False, mcp_config: Dict = None):
    async def llm_node_func(state: AgentState):
        print(">>> AGENT NODE: entered")
        messages = state['messages']
        llm = ChatGoogleGenerativeAI(
            google_api_key=os.getenv("GEMINI_API_KEY"), 
            model="gemini-2.5-flash",
            temperature=0,
            timeout=60,
        )
        
        all_tools = []
        if bind_tools:
            all_tools.append(tavily_tool)
            
        if mcp_config and mcp_config.get('command'):
            cmd = mcp_config['command']
            args = mcp_config.get('args', [])
            print(f">>> AGENT NODE: loading MCP tools from {cmd} {args}")
            mcp_tools = await get_mcp_tools(cmd, args)
            all_tools.extend(mcp_tools)
            
        if all_tools:
            print(f">>> AGENT NODE: binding {len(all_tools)} tools")
            llm = llm.bind_tools(all_tools)
            
        today = datetime.datetime.now().strftime("%B %d, %Y")
        final_system_msg = (
            f"{system_instruction}\n\nCurrent Date: {today}.\n"
            "IMPORTANT: When presenting your final answer, write in clean, well-structured markdown. "
            "Do NOT include raw JSON, tool output, or unformatted data in your response. "
            "Synthesize any information from tools into a polished, human-readable report."
        )
        final_messages = [SystemMessage(content=final_system_msg)] + messages
        
        try:
            print(f">>> AGENT NODE: calling LLM with {len(final_messages)} messages...")
            response = await llm.ainvoke(final_messages)
            has_tool_calls = bool(getattr(response, 'tool_calls', None))
            print(f">>> AGENT NODE: LLM responded. Has tool calls: {has_tool_calls}")
            return {"messages": [response]}
        except Exception as e:
            print(f"LLM Invocation Failed: {e}")
            return {"messages": [AIMessage(content=f"Error executing agent: {str(e)}")]}
        
    return llm_node_func

def get_email_node(receiver_email: str):
    def email_node_func(state: AgentState):
        print(f">>> EMAIL NODE: entered, sending to {receiver_email}")
        sender_email = os.getenv("EMAIL_USER")
        sender_password = os.getenv("EMAIL_PASS")

        if not sender_email or not sender_password:
            return {"messages": [AIMessage(content="Error: EMAIL_USER or EMAIL_PASS not found.")]}

        final_receiver = receiver_email
        if not final_receiver or not final_receiver.strip():
             final_receiver = sender_email
        
        last_message = state["messages"][-1]
        email_body = str(last_message.content)

        try:
            attachment_path = state.get("attachment_path", "")
        
            if attachment_path and os.path.isfile(attachment_path):
                # Send email WITH attachment
                msg = MIMEMultipart()
                msg['Subject'] = "Agent Report"
                msg['From'] = sender_email
                msg['To'] = final_receiver
                msg.attach(MIMEText(email_body, 'plain', 'utf-8'))
            
                with open(attachment_path, 'rb') as f:
                    part = MIMEBase('application', 'octet-stream')
                    part.set_payload(f.read())
                    encoders.encode_base64(part)
                    part.add_header(
                        'Content-Disposition',
                        f'attachment; filename="{os.path.basename(attachment_path)}"'
                    )
                    msg.attach(part)
            else:
                # Send plain-text email (no attachment)
                msg = MIMEText(email_body, 'plain', 'utf-8')
                msg['Subject'] = "Agent Report"
                msg['From'] = sender_email
                msg['To'] = final_receiver



            with smtplib.SMTP_SSL('smtp.gmail.com', 465) as server:
                server.login(sender_email, sender_password)
                server.send_message(msg)
            
            return {"messages": [AIMessage(content=f"✅ Email sent successfully to {final_receiver}")]}
        except Exception as e:
            return {"messages": [AIMessage(content=f"❌ Failed to send email: {str(e)}")]}
    return email_node_func

def get_whatsapp_node(receiver_phone: str):
    def whatsapp_node_func(state: AgentState):
        sid = os.getenv("TWILIO_ACCOUNT_SID")
        token = os.getenv("TWILIO_AUTH_TOKEN")
        from_number = os.getenv("TWILIO_FROM_NUMBER")

        if not sid or not token or not from_number:
            return {"messages": [AIMessage(content="Error: Twilio Credentials not found.")]}

        if not from_number.startswith("whatsapp:"):
            from_number = f"whatsapp:{from_number}"

        final_receiver = receiver_phone
        if not final_receiver or not final_receiver.strip():
             return {"messages": [AIMessage(content="Error: No Receiver Phone Number provided.")]}

        if not final_receiver.startswith("whatsapp:"):
            final_receiver = f"whatsapp:{final_receiver}"
            
        print(f"--- SENDING WHATSAPP TO {final_receiver} ---")

        last_message = state["messages"][-1]
        body_text = str(last_message.content)
        
        if len(body_text) > 1500:
            body_text = body_text[:1500] + "... (truncated)"

        try:
            client = Client(sid, token)
            message = client.messages.create(
                body=body_text,
                from_=from_number,
                to=final_receiver
            )
            return {"messages": [AIMessage(content=f"✅ WhatsApp sent! SID: {message.sid}")]}
        except TwilioRestException as e:
            if e.code == 63007:
                return {"messages": [AIMessage(content=f"❌ Failed: Twilio Channel not found (Sandbox not joined).")]}
            return {"messages": [AIMessage(content=f"❌ Failed to send WhatsApp: {str(e)}")]}
        except Exception as e:
            return {"messages": [AIMessage(content=f"❌ Failed to send WhatsApp: {str(e)}")]}

    return whatsapp_node_func

# --- UPDATED: DOC WRITER NODE (Word .docx) ---

import re

def _add_formatted_text(paragraph, text):
    """Parse inline markdown (bold **text**) and add runs to a paragraph."""
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)

def _unwrap_structured_content(raw_content: str) -> str:
    """
    Detect and unwrap structured/JSON content from LLM output.
    Handles: Gemini structured output, Tavily search results, etc.
    Returns clean text ready for markdown parsing.
    """
    content = raw_content.strip()

    # --- Try JSON parsing first ---
    try:
        import json
        parsed = json.loads(content)
        if isinstance(parsed, list) and len(parsed) > 0 and isinstance(parsed[0], dict):
            # LLM structured output: [{"type": "text", "text": "..."}]
            if all('text' in item for item in parsed):
                return '\n'.join(item['text'] for item in parsed)
            # Tavily search results: [{"title": "...", "url": "...", "content": "..."}]
            if all('content' in item for item in parsed):
                parts = []
                for item in parsed:
                    if item.get('title'):
                        parts.append(f"## {item['title']}")
                    if item.get('url'):
                        parts.append(f"*Source: {item['url']}*")
                    if item.get('content'):
                        parts.append(item['content'])
                    parts.append("")
                return '\n'.join(parts)
    except (json.JSONDecodeError, ValueError, TypeError):
        pass

    # --- Try Python literal eval (for single-quoted dicts) ---
    if content.startswith("[{") or content.startswith("({')"):
        try:
            import ast
            parsed = ast.literal_eval(content)
            if isinstance(parsed, list) and len(parsed) > 0 and isinstance(parsed[0], dict):
                if 'text' in parsed[0]:
                    return '\n'.join(item.get('text', '') for item in parsed)
                if 'content' in parsed[0]:
                    parts = []
                    for item in parsed:
                        if item.get('title'):
                            parts.append(f"## {item['title']}")
                        if item.get('url'):
                            parts.append(f"*Source: {item['url']}*")
                        if item.get('content'):
                            parts.append(item['content'])
                        parts.append("")
                    return '\n'.join(parts)
        except Exception:
            pass

    return content  # Return as-is if no structured format detected


def _parse_content_to_docx(doc, raw_content: str):
    """
    Parse markdown-style LLM output into formatted Word document elements.
    Handles: headings (#), bold (**), bullet points (* / -), and plain paragraphs.
    First unwraps any structured/JSON wrapper around the content.
    """
    content = _unwrap_structured_content(raw_content)

    # Split into lines (handle both literal \n and actual newlines)
    lines = content.replace('\\n', '\n').split('\n')

    for line in lines:
        stripped = line.strip()

        # Skip empty lines
        if not stripped:
            continue

        # --- Headings ---
        heading_match = re.match(r'^(#{1,4})\s+(.*)', stripped)
        if heading_match:
            level = len(heading_match.group(1))  # 1-4
            heading_text = heading_match.group(2).replace('**', '')
            doc.add_heading(heading_text, level=min(level, 4))
            continue

        # Skip lines with 5+ hashes (not valid Word heading levels)
        if re.match(r'^#{5,}\s+', stripped):
            text = re.sub(r'^#+\s+', '', stripped)
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.bold = True
            continue

        # --- Bullet points (* or -) ---
        bullet_match = re.match(r'^[\*\-]\s+(.*)', stripped)
        if bullet_match:
            bullet_text = bullet_match.group(1)
            p = doc.add_paragraph(style='List Bullet')
            _add_formatted_text(p, bullet_text)
            continue

        # --- Numbered list (1. / 2. etc.) ---
        numbered_match = re.match(r'^\d+\.\s+(.*)', stripped)
        if numbered_match:
            item_text = numbered_match.group(1)
            p = doc.add_paragraph(style='List Number')
            _add_formatted_text(p, item_text)
            continue

        # --- Regular paragraph ---
        p = doc.add_paragraph()
        _add_formatted_text(p, stripped)


def get_doc_writer_node(filename: str):
    """
    Writes content to a .docx file (compatible with Google Docs).
    Parses markdown-style LLM output into properly formatted Word elements.
    """
    def doc_writer_node_func(state: AgentState):
        print(">>> DOC WRITER NODE: entered")
        final_filename = filename if filename and filename.strip() else "agent_output.docx"
        if not final_filename.endswith(".docx"):
            final_filename += ".docx"
            
        # Use absolute path anchored to this script's directory
        script_dir = os.path.dirname(os.path.abspath(__file__))
        output_dir = os.path.join(script_dir, "generated_docs")
        os.makedirs(output_dir, exist_ok=True)
        
        file_path = os.path.join(output_dir, final_filename)
        print(f">>> DOC WRITER NODE: writing to {file_path}")
        
        last_message = state["messages"][-1]
        content = str(last_message.content)
        
        try:
            # Create a Word Document
            doc = DocxDocument()
            doc.add_heading('Agent Generated Report', 0)
            _parse_content_to_docx(doc, content)

            # Write to a temp file first, then rename — avoids
            # "Permission denied" when the target is open / locked.
            import tempfile
            tmp_fd, tmp_path = tempfile.mkstemp(suffix=".docx", dir=output_dir)
            os.close(tmp_fd)
            doc.save(tmp_path)
            # Replace the target file
            if os.path.exists(file_path):
                os.remove(file_path)
            shutil.move(tmp_path, file_path)
            
            abs_path = os.path.abspath(file_path)
            print(f">>> DOC WRITER NODE: success → {abs_path}")
            return {
                "messages": [AIMessage(content=f"✅ Word Document created: {abs_path}")],
                "attachment_path": abs_path
                }
        except Exception as e:
            print(f">>> DOC WRITER NODE: FAILED → {e}")
            return {"messages": [AIMessage(content=f"❌ Failed to write document: {str(e)}")]}

    return doc_writer_node_func

def get_google_sheets_node(spreadsheet_id: str, sheet_name: str):
    """Appends content from the agent to a Google Sheet."""
    def google_sheets_node_func(state: AgentState):
        print(f">>> GOOGLE SHEETS NODE: entered, sheet={sheet_name}")
        
        try:
            # Get user_id from metadata
            user_id = state.get("metadata", {}).get("user_id")
            print(f">>> GOOGLE SHEETS NODE: user_id from state: {user_id}")
            
            if not user_id:
                return {"messages": [AIMessage(content="❌ Error: user_id not found in state")]}
            
            # Import supabase from main (at function call time to avoid circular imports)
            from main import supabase, supabase_admin
            
            client = supabase_admin if supabase_admin else supabase
            
            if not client:
                return {"messages": [AIMessage(content="❌ Error: Supabase not initialized")]}
            
            # Fetch credentials from Supabase
            print(f">>> GOOGLE SHEETS NODE: querying user_integrations for user_id={user_id}")
            res = client.table("user_integrations").select("*").eq("user_id", user_id).eq("provider", "google").execute()
            print(f">>> GOOGLE SHEETS NODE: query result: {len(res.data) if res.data else 0} records found")
            
            if not res.data:
                return {"messages": [AIMessage(content="❌ Error: Google account not connected. Please sign in with Google.")]}
            
            token_info = res.data[0]
            
            # Create credentials object
            creds = Credentials(
                token=token_info["access_token"],
                refresh_token=token_info["refresh_token"],
                token_uri="https://oauth2.googleapis.com/token",
                client_id=os.getenv("GOOGLE_CLIENT_ID"),
                client_secret=os.getenv("GOOGLE_CLIENT_SECRET")
            )
            
            # Refresh credentials if expired
            if creds.expired and creds.refresh_token:
                creds.refresh(Request())
                # Update the new access token in Supabase
                client.table("user_integrations").update({"access_token": creds.token}).eq("id", token_info["id"]).execute()
            
            # Connect to Google Sheets
            print(f">>> GOOGLE SHEETS NODE: spreadsheet_id={spreadsheet_id}, sheet_name={sheet_name}")
            if not spreadsheet_id:
                return {"messages": [AIMessage(content="❌ Error: No spreadsheet ID configured. Please set the Spreadsheet ID in the node settings.")]}
            
            client = gspread.authorize(creds)
            sheet = client.open_by_key(spreadsheet_id).worksheet(sheet_name)
            
            # Get content from the last message
            last_message = state["messages"][-1]
            content = str(last_message.content)
            
            # Append data to sheet
            sheet.append_row([content])
            
            return {"messages": [AIMessage(content=f"✅ Successfully appended to Google Sheet: {sheet_name}")]}
            
        except Exception as e:
            print(f">>> GOOGLE SHEETS NODE: FAILED → {e}")
            return {"messages": [AIMessage(content=f"❌ Google Sheets Error: {str(e)}")]}
    
    return google_sheets_node_func

# --- 5. GRAPH BUILDER ---

async def build_and_run_workflow(nodes_config: List[Dict], edges_config: List[Dict], request_initial_input: str, user_id: str = None):
    print(f"Building workflow with {len(nodes_config)} nodes")

    workflow = StateGraph(AgentState)
    
    has_native_tools = any(n.get('data', {}).get('backendType') in ['tool', 'search'] for n in nodes_config)
    
    mcp_config = None
    mcp_node = next((n for n in nodes_config if n.get('data', {}).get('backendType') == 'mcp'), None)
    if mcp_node:
        raw_cmd = mcp_node['data'].get('serverCommand', '')
        if raw_cmd:
            parts = raw_cmd.split(' ')
            mcp_config = {"command": parts[0], "args": parts[1:]}

    # 1. Add Nodes
    input_override = ""

    for node in nodes_config:
        node_id = node['id']
        data = node.get('data', {})
        backend_type = data.get('backendType', 'default')
        
        print(f">>> Creating node: {node_id}, backendType: {backend_type}")

        if backend_type == 'input':
            input_override = data.get('userPrompt', '')
            workflow.add_node(node_id, lambda state: {}) 

        elif backend_type == 'agent':
            sys_instr = data.get('systemInstruction', 'You are a helpful assistant.')
            user_tmpl = data.get('promptTemplate', '{input}')
            workflow.add_node(node_id, get_llm_node(sys_instr, user_tmpl, bind_tools=has_native_tools, mcp_config=mcp_config))
            
        elif backend_type == 'tool' or backend_type == 'search':
            workflow.add_node(node_id, ToolNode([tavily_tool]))
            
        elif backend_type == 'mcp':
            workflow.add_node(node_id, lambda state: {})
            
        elif backend_type == 'email':
            raw_receiver = data.get('receiverEmail')
            receiver = raw_receiver if raw_receiver and raw_receiver.strip() else os.getenv("RECEIVER_EMAIL") or os.getenv("EMAIL_USER")
            workflow.add_node(node_id, get_email_node(receiver))
            
        elif backend_type == 'whatsapp':
            receiver = data.get('receiverPhone')
            workflow.add_node(node_id, get_whatsapp_node(receiver))
            
        elif backend_type == 'doc_writer':
            filename = data.get('filename')
            workflow.add_node(node_id, get_doc_writer_node(filename))
            
        elif backend_type == 'google_sheets':
            spreadsheet_id = data.get('spreadsheetId')
            sheet_name = data.get('sheetName', 'Sheet1')
            print(f">>> Creating Google Sheets node: spreadsheet_id={spreadsheet_id}, sheet_name={sheet_name}")
            workflow.add_node(node_id, get_google_sheets_node(spreadsheet_id, sheet_name))
        
        else:
            # Fallback for unknown types - add a passthrough node
            print(f"⚠️  Unknown node type '{backend_type}' for node {node_id}. Adding passthrough node.")
            workflow.add_node(node_id, lambda state: {})
    
    # Get all valid node IDs
    valid_node_ids = {node['id'] for node in nodes_config}
    print(f">>> Valid node IDs: {valid_node_ids}")
    
    # Filter edges to only include those with valid source and target nodes
    edges_config_filtered = [
        edge for edge in edges_config
        if edge.get('source') in valid_node_ids and edge.get('target') in valid_node_ids
    ]
    
    if len(edges_config_filtered) < len(edges_config):
        print(f"⚠️  Filtered out {len(edges_config) - len(edges_config_filtered)} invalid edges")
    
    # 2. Add Edges — with proper ReAct loop for Agent ↔ Tool cycling
    #    Before: Agent → Tool → DocWriter  (tool's raw JSON goes straight to doc writer)
    #    After:  Agent ↔ Tool (loop), then Agent → DocWriter when done with tools

    # First pass: identify agent→tool pairs and tool→next_node successors
    agent_to_tool = {}   # agent_id -> tool_id
    tool_successor = {}  # tool_id -> next_node_id

    for edge in edges_config_filtered:
        src = edge['source']
        tgt = edge['target']
        src_node = next((n for n in nodes_config if n['id'] == src), None)
        tgt_node = next((n for n in nodes_config if n['id'] == tgt), None)
        src_type = src_node.get('data', {}).get('backendType') if src_node else None
        tgt_type = tgt_node.get('data', {}).get('backendType') if tgt_node else None

        if src_type == 'agent' and tgt_type in ('tool', 'search', 'mcp'):
            agent_to_tool[src] = tgt
        if src_type in ('tool', 'search', 'mcp'):
            tool_successor[src] = tgt

    # Second pass: build edges
    handled_tool_edges = set()  # tool→next edges replaced by the ReAct loop

    for edge in edges_config_filtered:
        source = edge['source']
        target = edge['target']
        source_node = next((n for n in nodes_config if n['id'] == source), None)
        target_node = next((n for n in nodes_config if n['id'] == target), None)
        source_type = source_node.get('data', {}).get('backendType') if source_node else None
        target_type = target_node.get('data', {}).get('backendType') if target_node else None

        if source_type == 'agent' and target_type in ('tool', 'search', 'mcp'):
            # --- ReAct pattern: Agent ↔ Tool, then Agent → next node ---
            next_after_tool = tool_successor.get(target)
            if next_after_tool:
                workflow.add_conditional_edges(
                    source, tools_condition, {"tools": target, END: next_after_tool}
                )
                handled_tool_edges.add((target, next_after_tool))
            else:
                workflow.add_conditional_edges(
                    source, tools_condition, {"tools": target, END: END}
                )
            # Loop tool results back to agent for synthesis
            workflow.add_edge(target, source)
            print(f"🔄 ReAct loop: {source} ↔ {target}, exit → {next_after_tool or 'END'}")

        elif (source, target) in handled_tool_edges:
            # Skip — this edge is now handled by the ReAct loop's conditional routing
            continue

        else:
            workflow.add_edge(source, target)
            
    # 3. Entry Point & Run
    start_node = next((n['id'] for n in nodes_config if n.get('data', {}).get('backendType') == 'input'), None)
    if start_node: workflow.set_entry_point(start_node)
    
    graph_app = workflow.compile()
    # User's chat message takes priority; input node prompt is a fallback for test runs
    final_input = request_initial_input if request_initial_input.strip() else input_override

    print(f">>> WORKFLOW: starting execution with input: {final_input[:100]}...")

    try:
        final_state = await asyncio.wait_for(
            graph_app.ainvoke({
                "messages": [HumanMessage(content=final_input)],
                "metadata": {"user_id": user_id} if user_id else {},
                "attachment_path": ""
            }),
            timeout=180  # 3 minute timeout
        )
        print(f">>> WORKFLOW: execution completed successfully")
        return {
            "result": final_state["messages"][-1].content,
            "full_history": [m.content for m in final_state["messages"]]
        }
    except asyncio.TimeoutError:
        print(f">>> WORKFLOW: TIMED OUT after 180s")
        return {
            "result": "Error: Workflow timed out after 3 minutes.",
            "full_history": []
        }
    except Exception as e:
        print(f"Workflow execution failed: {e}")
        return {
            "result": f"Error: {str(e)}",
            "full_history": []
        }